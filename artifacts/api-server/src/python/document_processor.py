#!/usr/bin/env python3
"""
OMNIMENS Document Processor — Word (.docx) + Excel (.xlsx)
STDIN: JSON {action, file_b64, options/content}
Actions: read_docx, create_docx, read_excel, create_excel, analyze_csv
"""
import sys, json, base64, io, tempfile, os
def error_out(msg): print(json.dumps({"success": False, "error": msg})); sys.exit(0)

def process(spec: dict) -> dict:
    action = spec.get("action", "read_docx")
    file_b64 = spec.get("file_b64", "")
    file_bytes = None
    if file_b64:
        try: file_bytes = base64.b64decode(file_b64)
        except: error_out("Invalid base64 file")

    # ── Word Document ────────────────────────────────────────────────────────
    if action == "read_docx":
        from docx import Document
        if not file_bytes: error_out("file_b64 required")
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [{"index": i, "text": p.text, "style": p.style.name}
                      for i, p in enumerate(doc.paragraphs) if p.text.strip()]
        tables = []
        for ti, table in enumerate(doc.tables):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append({"table_index": ti, "rows": rows, "row_count": len(rows)})
        full_text = "\n".join(p["text"] for p in paragraphs)
        return {"success": True, "action": "read_docx", "paragraph_count": len(paragraphs),
                "table_count": len(tables), "paragraphs": paragraphs, "tables": tables,
                "full_text": full_text[:30000]}

    elif action == "create_docx":
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        content = spec.get("content", "")
        title = spec.get("title", "OMNIMENS Document")
        sections = spec.get("sections", [])
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)
        # Title
        h = doc.add_heading(title, 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if content:
            for line in content.split('\n'):
                line = line.strip()
                if line.startswith("# "): doc.add_heading(line[2:], 1)
                elif line.startswith("## "): doc.add_heading(line[3:], 2)
                elif line.startswith("### "): doc.add_heading(line[4:], 3)
                elif line: doc.add_paragraph(line)
        for section in sections:
            if section.get("heading"): doc.add_heading(section["heading"], 2)
            if section.get("body"): doc.add_paragraph(section["body"])
            if section.get("table"):
                tdata = section["table"]
                if tdata:
                    t = doc.add_table(rows=len(tdata), cols=len(tdata[0]))
                    t.style = "Table Grid"
                    for i, row in enumerate(tdata):
                        for j, cell_text in enumerate(row):
                            t.cell(i, j).text = str(cell_text)
        buf = io.BytesIO(); doc.save(buf)
        return {"success": True, "action": "create_docx",
                "docx_base64": base64.b64encode(buf.getvalue()).decode(),
                "size_bytes": len(buf.getvalue())}

    # ── Excel Spreadsheet ────────────────────────────────────────────────────
    elif action == "read_excel":
        import openpyxl
        if not file_bytes: error_out("file_b64 required")
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        sheets = {}
        for name in wb.sheetnames:
            ws = wb[name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            # Trim empty trailing rows
            while rows and all(c == "" for c in rows[-1]): rows.pop()
            sheets[name] = {"rows": rows, "row_count": len(rows),
                            "col_count": max((len(r) for r in rows), default=0)}
        return {"success": True, "action": "read_excel", "sheet_names": wb.sheetnames, "sheets": sheets}

    elif action == "create_excel":
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        title = spec.get("title", "OMNIMENS Spreadsheet")
        sheets_data = spec.get("sheets", [])
        if not sheets_data and spec.get("data"): sheets_data = [{"name": "Sheet1", "data": spec["data"]}]
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        HEADER_FILL = PatternFill("solid", fgColor="6C63FF")
        HEADER_FONT = Font(bold=True, color="FFFFFF")
        EVEN_FILL = PatternFill("solid", fgColor="F0EFFF")
        for si, sheet_spec in enumerate(sheets_data):
            ws = wb.create_sheet(sheet_spec.get("name", f"Sheet{si+1}"))
            data = sheet_spec.get("data", [])
            for ri, row in enumerate(data):
                for ci, val in enumerate(row):
                    cell = ws.cell(row=ri+1, column=ci+1, value=val)
                    if ri == 0:
                        cell.font = HEADER_FONT; cell.fill = HEADER_FILL
                        cell.alignment = Alignment(horizontal="center")
                    elif ri % 2 == 0: cell.fill = EVEN_FILL
            for col in ws.columns:
                max_len = max((len(str(c.value or "")) for c in col), default=0)
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
        buf = io.BytesIO(); wb.save(buf)
        return {"success": True, "action": "create_excel",
                "excel_base64": base64.b64encode(buf.getvalue()).decode(),
                "size_bytes": len(buf.getvalue())}

    elif action == "analyze_csv":
        import csv, io as sysio
        import pandas as pd
        content = spec.get("content", "")
        if not content and file_bytes: content = file_bytes.decode("utf-8", errors="replace")
        if not content: error_out("content or file_b64 required")
        df = pd.read_csv(sysio.StringIO(content))
        describe = df.describe(include="all").fillna("").to_dict()
        return {"success": True, "action": "analyze_csv",
                "shape": list(df.shape), "columns": list(df.columns),
                "dtypes": {c: str(t) for c, t in df.dtypes.items()},
                "describe": describe, "head": df.head(10).to_dict(orient="records"),
                "null_counts": df.isnull().sum().to_dict(),
                "value_counts": {c: df[c].value_counts().head(5).to_dict()
                                 for c in df.select_dtypes(include="object").columns[:5]}}

    else:
        error_out(f"Unknown action: {action}. Use: read_docx, create_docx, read_excel, create_excel, analyze_csv")

if __name__ == "__main__":
    raw = sys.stdin.read().strip()
    if not raw: error_out("No input")
    try: spec = json.loads(raw)
    except: error_out("Invalid JSON")
    print(json.dumps(process(spec)))
